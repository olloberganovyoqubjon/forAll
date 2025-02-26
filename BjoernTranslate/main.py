import os
import socket
from flask import Flask, request, send_file, jsonify
from docx import Document
from io import BytesIO
import transliterate
from flask_cors import CORS
import win32com.client as win32
from werkzeug.utils import secure_filename
import pythoncom
from py_eureka_client import eureka_client
from dotenv import load_dotenv
import tempfile

load_dotenv()

app = Flask(__name__)

@app.route("/text", methods=['POST'])
def textTranslate():
    if request.method == 'POST':
        print(request.date)
        request_data = request.get_json()
        lang = request_data['lang']
        text = request_data['text']
        print("textIn: ", text)
        text = transliterate.transliterate(text, lang)
        print("textEx: ", text)
    return jsonify(text)

@app.route('/upload/<lang>', methods=['POST'])
def upload_file(lang):
    if 'file' not in request.files:
        return 'No file part', 400
    file = request.files['file']
    if file.filename == '':
        return 'No selected file', 400

    file_stream = BytesIO(file.read())  # Faylni xotirada o'qib olish

    if file.filename.endswith('.doc'):
        translated_doc_bytes = translate_doc_file(file_stream, lang)
        if translated_doc_bytes:
            return send_file(
                translated_doc_bytes,
                mimetype="application/msword",
                download_name="translated_document.doc"
            )
        else:
            return 'Failed to process .doc file', 500

    elif file.filename.endswith('.docx'):
        translated_docx_bytes = translate_docx_file(file_stream, lang)
        return send_file(translated_docx_bytes, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", download_name="translated_document.docx")

    return 'Invalid file type', 400

def translate_doc_file(file_stream, lang):
    """ .doc faylni xotirada saqlab, tarjima qilib, .doc formatida qaytaradi """
    try:
        pythoncom.CoInitialize()
        word = win32.Dispatch("Word.Application")
        word.Visible = False  # GUI'ni yashirish

        # 🛠 1. Hujjatni ochish
        doc = word.Documents.Open(file_stream)

        # 🛠 2. Har bir paragrafni transliteratsiya qilish
        for paragraph in doc.Paragraphs:
            paragraph.Range.Text = transliterate.transliterate(paragraph.Range.Text, lang)

        # 🛠 3. `.doc` formatida xotirada (`BytesIO`) saqlash
        translated_doc_bytes = BytesIO()
        temp_output_path = os.path.join(app.config['UPLOAD_FOLDER'], "temp_translated.doc")
        doc.SaveAs(temp_output_path, FileFormat=0)  # `.doc` formatida saqlash

        doc.Close(SaveChanges=False)
        word.Quit()

        with open(temp_output_path, "rb") as f:
            translated_doc_bytes.write(f.read())

        os.remove(temp_output_path)  # `.doc` faylni o‘chirish

        translated_doc_bytes.seek(0)
        return translated_doc_bytes

    except Exception as e:
        print(f"Error processing .doc file: {e}")
        return None
    finally:
        pythoncom.CoUninitialize()

def translate_docx_file(file_stream, lang):
    """ .docx faylni transliterate qilib, qaytaradi """
    try:
        translated_docx_bytes = BytesIO()
        doc = Document(file_stream)

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.text = transliterate.transliterate(run.text, lang)

        doc.save(translated_docx_bytes)
        translated_docx_bytes.seek(0)
        return translated_docx_bytes

    except Exception as e:
        print(f"Error processing .docx file: {e}")
        return None

def translate_paragraph(paragraph, lang):
    for run in paragraph.runs:
        run.text = transliterate.transliterate(run.text, lang)

def translate_document(doc, lang):
    for paragraph in doc.paragraphs:
        translate_paragraph(paragraph, lang)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    translate_paragraph(paragraph, lang)

EUREKA_SERVER = "http://localhost:8761/eureka"
YOUR_APP_NAME = "python_translate"
YOUR_INSTANCE_PORT = 5000  # Flask default port

def start_eureka_client():
    eureka_client.init(eureka_server=EUREKA_SERVER,
                       app_name=YOUR_APP_NAME,
                       instance_port=YOUR_INSTANCE_PORT)

if __name__ == '__main__':
    start_eureka_client()
    app.run(debug=True)
